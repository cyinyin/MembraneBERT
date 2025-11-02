import os
import re
from docx import Document


# Read text from a Word document
def read_word_file(file_path):
    doc = Document(file_path)
    text = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    return text


# Tokenization rules: handle scientific notation, compound units, and punctuation
def split_text_with_tags(line):
    pattern = r"""
        \d+\.\d+(?:×10\^[-+]?\d+)?%?        |  # Decimals or scientific notation with an optional percent sign, e.g., 3.14, 3.14×10^-5%
        \d+%                                 |  # Integer with a percent sign, e.g., 98%
        \w+%                                 |  # # Word with a percent sign, e.g., wt%
        \[.*?\]                        |  # Content within square brackets treated as a whole, e.g., [emim]
        \b(?:\w+(?:[-/]\w+)*)(?:/?)\b        |  # Words connected by hyphens or slashes, e.g., PSF/SAPO-34/
        \w+(?:²|³)?(?:/\w+)*                 |  # Compound units, e.g., cm²/s
        [^\w\s]                              # Punctuation
    """
    words = re.findall(pattern, line, re.VERBOSE)
    return words, ['O'] * len(words)


def write_labeled_to_word(text, tags, output_path):
    doc = Document()
    for i, line in enumerate(text):
        words, _ = split_text_with_tags(line)
        labeled_line = []
        for j, word in enumerate(words):
            labeled_line.append(f"{word}: {tags[i][j]}")  # Add the label O after each token
        doc.add_paragraph(' '.join(labeled_line))
    doc.save(output_path)


def process_all_files(input_dir, output_dir):

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for file_name in os.listdir(input_dir):
        if file_name.endswith(".docx"):
            input_file = os.path.join(input_dir, file_name)
            output_file = os.path.join(output_dir, file_name)

            text = read_word_file(input_file)

            tags = []
            for line in text:
                _, line_tags = split_text_with_tags(line)
                tags.append(line_tags)

            write_labeled_to_word(text, tags, output_file)
            print(f"Processed {file_name}")

input_directory = "D:/organic_frameworks_extracts"
output_directory = "D:/tag_organic_frameworks_text"

# Process all Word files
process_all_files(input_directory, output_directory)
