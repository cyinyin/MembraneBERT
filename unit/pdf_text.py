import fitz  # PyMuPDF
import os
import re
from docx import Document
from docx.shared import Pt
from collections import Counter
import threading
from unit.path import Args


source_folder = Args().pdf_path
output_folder = Args().plain_text_path

if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# Get all PDF files
pdf_files = [os.path.join(source_folder, f) for f in os.listdir(source_folder) if f.endswith('.pdf')]


# Timeout exception
class TimeoutException(Exception):
    pass


def timeout_handler():
    raise TimeoutException("File read timeout")


# 从PDF提取文本
def extract_text_from_pdf(pdf_path):
    pdf_document = fitz.open(pdf_path)
    text = ""
    for page in pdf_document:
        page_text = page.get_text("text")
        cleaned_text = clean_text(page_text)
        text += cleaned_text + "\n"
    pdf_document.close()
    return text


def clean_scientific_notation(text):
    """
    Fix issues with scientific notation extraction, including:

    Correction of misaligned standard scientific notation.
    Fix for missing multiplication signs in consecutive multiplication patterns.
    """
    # Fix standard scientific notation
    pattern_scientific = re.compile(
        r'(\b\d+(\.\d+)?)[\s×x]*(\d+)?[\s×x]*10[\s\^]*([-+]?\d+)', re.IGNORECASE
    )

    def replace_scientific(match):
        """
        Correct the format of scientific notation.
        """
        base = match.group(1)
        multiplier = match.group(3) if match.group(3) else ""
        exponent = match.group(4)

        if multiplier:
            base = f"{base}{multiplier}"

        return f"{base}×10^{exponent}"

    text = pattern_scientific.sub(replace_scientific, text)

    pattern_multiplication = re.compile(
        r'(\b\d+(\.\d+)?)\s+(\d+(\.\d+)?×10\^[-+]?\d+)', re.IGNORECASE
    )

    def replace_multiplication(match):
        """
        Insert missing multiplication signs in consecutive multiplication patterns.
        """
        first_number = match.group(1)
        second_part = match.group(3)
        return f"{first_number}×{second_part}"

    text = pattern_multiplication.sub(replace_multiplication, text)
    return text


def clean_directory_lines(text):
    text = re.sub(r'\.{2,}', '......', text)
    text = re.sub(r'[-–—]{2,}', '......', text)
    text = re.sub(r'\s*\.\s*', '.', text)
    text = re.sub(r'\s*\-\s*', '-', text)
    return text


def normalize_directory_format(text):
    pattern = re.compile(r'(.*?)\s+(\.{3,}|[-—]{3,})\s+(\d+)', re.IGNORECASE)
    return pattern.sub(r'\1 ...... \3', text)


def clean_text(text):
    lines = text.splitlines()
    cleaned_lines = []
    in_table = False
    short_line_count = 0

    for line in lines:
        if re.match(r'^\s*[-—+|]*\s*$', line):
            continue
        line = clean_directory_lines(line)
        line = normalize_directory_format(line)
        line = re.sub(r'\[\d+([–-]?\d+)*(\s*,\s*\d+)*\]', '', line)
        line = re.sub(r'\(\d+([–-]?\d+)*\)', '', line)

        line = re.sub(r'(\d)\s*\.(\s*\d)', r'\1.\2', line)

        line = re.sub(r'1\s*C', '℃', line)
        line = re.sub(r'(\d+)\s*[°]?\s*C', r'\1℃', line)

        if len(line.strip()) < 20:
            short_line_count += 1
        else:
            short_line_count = 0

        if short_line_count >= 3:
            in_table = True

        if in_table and short_line_count == 0:
            in_table = False

        if not in_table:
            cleaned_lines.append(line)

    return "\n".join(cleaned_lines)


# Clean up headers and footers
def clean_headers_and_footers(text):
    lines = text.splitlines()
    line_counter = Counter(lines)
    cleaned_lines = []
    for line in lines:
        if line_counter[line] > 2 and len(line) < 100:
            continue
        if re.match(r'^\s*(Page\s+\d+|DOI|Copyright|Received|Accepted|Elsevier|Springer)', line, re.IGNORECASE):
            continue
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines)


# Clean up incompatible characters
def clean_invalid_xml_chars(text):
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x84\x86-\x9F]', '', text)


# Set the font to Times New Roman
def set_times_new_roman_style(paragraph):
    run = paragraph.add_run()
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    return run


# Split the text into paragraphs and add them to the Word document
def add_text_with_line_breaks(doc, text):
    lines = text.splitlines()
    combined_line = ""
    paragraph = doc.add_paragraph()

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        combined_line += line + " "

        ends_with_punctuation = re.search(r'[.!?]$', line)
        next_line_starts_with_capital = (
            i + 1 < len(lines) and re.match(r'^[A-Z]', lines[i + 1])
        )

        if ends_with_punctuation and next_line_starts_with_capital:
            set_times_new_roman_style(paragraph)
            paragraph.add_run(combined_line.strip())
            paragraph = doc.add_paragraph()
            combined_line = ""

    if combined_line:
        set_times_new_roman_style(paragraph)
        paragraph.add_run(combined_line.strip())


# Remove the content of the References section
def remove_references(text):
    """
    Delete References and all subsequent content from the text.
    """
    ref_pattern = re.compile(r'^\s*References[:\s]*$', re.IGNORECASE)

    paragraphs = text.splitlines()
    remove_from_index = None

    for i in range(len(paragraphs) - 1, -1, -1):
        if ref_pattern.fullmatch(paragraphs[i].strip()):
            remove_from_index = i
            break

    if remove_from_index is not None:
        print(f"Removing content starting from: {paragraphs[remove_from_index].strip()}")
        return "\n".join(paragraphs[:remove_from_index])


    print("No References section found.")
    return text


def remove_acknowledgements(text):
    """
    Remove the Acknowledgements or Acknowledgment section and all subsequent content.
    """
    # Match the Acknowledgements or Acknowledgment heading and its following content
    ack_patterns = r'(?i)\bAcknowledgements?\b[:\s]*.*'  # Match heading and following content, case-insensitive

    # Remove everything from the matched heading to the end of the text
    match = re.search(ack_patterns, text, re.DOTALL)  # Use re.DOTALL to include newlines in the match

    if match:
        print(f"Removing Acknowledgements starting from: {match.group(0)[:50]}...")
        return text[:match.start()].strip()

    print("No Acknowledgements section found.")
    return text


def remove_contents(text):
    # Regex 1: Match table of contents entries starting with section numbers followed by titles and page numbers
    pattern1 = r'(?:\n*\d+\.\s+[^\n]*\s*\.\s*\d+\n?)+'
    # Regex 2: Match table of contents entries with titles followed by dots and page numbers
    pattern2 = r'\n*[^\n]+\.\.\.+(?:\s+\d+\.)?\s*(?:\d+\.)?\n?'
    # Regex 3: Match sequences of section numbers possibly followed by dots and numbers
    pattern3 = r'(?:\d+\.)+(?:\s*\.\s*)+(?:\d+\.)+(?:\s*\d+\.)*'

    # Remove matched table of contents sections
    text = re.sub(pattern1, '', text, flags=re.MULTILINE)
    text = re.sub(pattern2, '', text, flags=re.MULTILINE)
    text = re.sub(pattern3, '', text, flags=re.MULTILINE)

    return text.strip()


# Remove figure and table captions
def remove_figure_and_table_captions(text):
    text = re.sub(r'^(Figure|Fig\.|Figs\.)\s?\d+.*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'^(Table)\s?\d+.*', '', text, flags=re.IGNORECASE)
    return text


def process_multiple_pdfs(pdf_files, output_folder, timeout_seconds):
    for pdf_file in pdf_files:
        timer = threading.Timer(timeout_seconds, timeout_handler)
        try:
            print(f"Processing: {pdf_file}")
            timer.start()
            text = extract_text_from_pdf(pdf_file)
            if not text.strip():
                print(f"No extractable text found in {pdf_file}")
                continue

            text = clean_headers_and_footers(text)
            text = clean_invalid_xml_chars(text)
            text = clean_scientific_notation(text)
            text = remove_figure_and_table_captions(text)
            text = remove_contents(text)
            text = remove_references(text)
            text = remove_acknowledgements(text)

            doc = Document()
            add_text_with_line_breaks(doc, text)
            output_path = os.path.join(output_folder, os.path.splitext(os.path.basename(pdf_file))[0] + ".docx")
            doc.save(output_path)
            print(f"Saved: {output_path}")

        except TimeoutException:
            print(f"File {pdf_file} read timeout, skipped.")
            continue
        except Exception as e:
            print(f"Error processing file {pdf_file}: {e}")
            continue
        finally:
            timer.cancel()


# Main function call
process_multiple_pdfs(pdf_files, output_folder, timeout_seconds=60)
